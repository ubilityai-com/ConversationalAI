from docx.oxml import OxmlElement
from docx import Document
from io import BytesIO
import aiohttp, sys, os, json
import base64
sys.path.append(os.path.dirname(os.path.abspath(__file__)))
from applications.functions import get_file_data, upload_file


status = [200, 201, 202, 204, 206, 207, 208]

###################################### Replace Strings #################################################
async def replace_text_in_paragraph(paragraph, old_text, new_text):
    """Replace text in a paragraph while preserving formatting."""
    for run in paragraph.runs:
        if old_text in run.text:
            formatted_text = run.text.replace(old_text, str(new_text))
            run.text = str(formatted_text)

async def replace_text_in_cell(cell, old_text, new_text):
    """Replace text in a table cell while preserving formatting."""
    for paragraph in cell.paragraphs:
        replace_text_in_paragraph(paragraph, old_text, new_text)

async def copy_cell_formatting(old_cell, new_cell):
    # Copy cell font
    for old_paragraph, new_paragraph in zip(old_cell.paragraphs, new_cell.paragraphs):
        for old_run, new_run in zip(old_paragraph.runs, new_paragraph.runs):
            new_run.font.size = old_run.font.size
            new_run.font.bold = old_run.font.bold
            new_run.font.italic = old_run.font.italic
            new_run.font.underline = old_run.font.underline
            new_run.font.color.rgb = old_run.font.color.rgb
            new_run.font.name = old_run.font.name

    # Copy cell shading (background color) via XML
    shd_elements = old_cell._element.xpath('.//w:shd')
    if shd_elements:
        old_shd = shd_elements[0]
        new_shd = OxmlElement('w:shd')
        for attr in old_shd.attrib:
            new_shd.set(attr, old_shd.get(attr))
        new_cell._element.get_or_add_tcPr().append(new_shd)

async def microsoft_word_replace_strings(cred,params,**kwargs):
    """
    in the entered document, using a dictionary containing old strings and new strings as reference. replaces the old strings in the document with the new strings then returns the modified document.

    :accessToken: String Access token for authentication with Microsoft Graph API.
    :params: Dictionary containing parameters.

    - :values: (dict, Required) - dictionay containing the old strings to be replaced as keys and new strings to replace the old strings as values
    - :base64_string: (str) - the name used by the platform to identify the document. Required if 'direct_link' is undefined
    - :direct_link: (str) - direct download link for the document. Required if 'base64_string' is undefined
    
    Returns:
        dict: contains a confirmation message as well as the base64 content of the new document
    """
    try:
        if "values" in params :
            values = params["values"]
            old_doc = None
            
            if "base64_string" in params:
                base64_string = params["base64_string"]
                content = base64_string
                # Retrieve file content data
                if kwargs:
                    # Extra conv_id & dialogue_id
                    dialogue_id = kwargs.get("dialogue_id")
                    conv_id = kwargs.get("conv_id")
                contentData = get_file_data(dialogue_id,conv_id,content)
                if "Error" in contentData:
                    raise Exception(f"Failed to retrieve file content: {contentData['Error']}")
                # Extract and decode the file content from hex to bytes
                doc_content = bytes.fromhex(contentData["file_content"])
                old_doc = BytesIO(doc_content)
                old_doc.seek(0)
                original_filename = contentData["file_name"]
            elif "direct_link"in params:
                direct_link = params["direct_link"]
                # derive a filename from the URL
                original_filename = direct_link.split("/")[-1]
                async with aiohttp.ClientSession() as session:
                    async with session.get(direct_link) as response:
                        response.raise_for_status()
                        if response.status in status:
                            body = await response.read()
                            old_doc = BytesIO(body)
                            old_doc.seek(0)
                        else:
                            raise Exception(f"Status code: {response.status}. Response: {response.text}")
            else:
                raise Exception("missing input data")
            
            doc = Document(old_doc)

            # replace text in paragraphs
            for paragraph in doc.paragraphs:
                for (old_string, new_string) in values.items():
                    if old_string in paragraph.text:
                        replace_text_in_paragraph(paragraph, old_string, new_string) 
            
            # replace text in cells
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for (old_string, new_string) in values.items():
                            if old_string in cell.text:
                                replace_text_in_cell(cell, old_string, new_string)
            doc_new = BytesIO()
            doc.save(doc_new)
            doc_new.seek(0)
            
            
            new_doc = doc_new.read()
            if kwargs:
                # Extra conv_id & dialogue_id
                dialogue_id = kwargs.get("dialogue_id")
                conv_id = kwargs.get("conv_id")
                fileData = upload_file(dialogue_id, conv_id, new_doc,original_filename)
            return fileData
            
        else:
            raise Exception("missing input data")
    except aiohttp.ClientError as e:
        return {"Error": str(e)}
    except Exception as err:
        return {"Error": str(err)}


async def microsoft_word_replace_table(cred,params,**kwargs):
    try:
        if "values" in params and "index" in params :
            values = params["values"]
            index = params["index"]
            old_doc = None
            
            if "base64_string" in params :
                content = params["base64_string"]
                # Retrieve file content data
                if kwargs:
                    # Extra conv_id & dialogue_id
                    dialogue_id = kwargs.get("dialogue_id")
                    conv_id = kwargs.get("conv_id")
                contentData = get_file_data(dialogue_id,conv_id,content)
                if "Error" in contentData:
                    raise Exception(f"Failed to retrieve file content: {contentData['Error']}")
                # Extract and decode the file content from hex to bytes
                doc_content = bytes.fromhex(contentData["file_content"])
                old_doc = BytesIO(doc_content)
                old_doc.seek(0)
                original_filename = contentData["file_name"]
            elif "direct_link" in params :
                direct_link = params["direct_link"]
                # derive a filename from the URL
                original_filename = direct_link.split("/")[-1]
                async with aiohttp.ClientSession() as session:
                    async with session.get(direct_link) as response:
                        response.raise_for_status()
                        if response.status in status:
                            body = await response.read()
                            old_doc = BytesIO(body)
                            old_doc.seek(0)
                        else:
                            raise Exception(
                                f"Status code: {response.status}. Response: {response.text}")
            else:
                raise Exception("missing input data")
            
            doc = Document(old_doc)
            old_table = doc.tables[index]

            # Step 1: Create the new table with one row (header)
            new_table = doc.add_table(rows=1, cols=len(values[0]), style='Table Grid')
            header_values = values[0]

            # Step 2: Set header values and format
            header = new_table.rows[0]
            for idx, new_cell in enumerate(header.cells):
                # Add header text to the cell
                paragraph = new_cell.paragraphs[0]
                run = paragraph.add_run(header_values[idx])  # Add the corresponding header value
                old_cell = old_table.rows[0].cells[idx]  # Assuming the first row is the header row
                copy_cell_formatting(old_cell, new_cell) # Add the corresponding header format
            
            # Step 3: Set other cells values
            del values[0] # Remove header values
            for row_data in values:
                # Add a new row to the table
                row = new_table.add_row().cells
                for idx, value in enumerate(row_data):
                    row[idx].text = str(value)  # Set the text


            # Step 4: Insert the new table at the same position as the old table
            old_table_element = old_table._element
            old_table_element.addnext(new_table._element)

            # Step 5: Remove the old table
            old_table_element.getparent().remove(old_table_element)  

            doc_new = BytesIO()
            doc.save(doc_new)
            doc_new.seek(0)
            
            
            new_doc = doc_new.read()
            if kwargs:
                # Extra conv_id & dialogue_id
                dialogue_id = kwargs.get("dialogue_id")
                conv_id = kwargs.get("conv_id")
                fileData = upload_file(dialogue_id, conv_id, new_doc,original_filename)
            return fileData

    except aiohttp.ClientError as e:
        return {"Error": str(e)}
    except Exception as err:
        return {"Error": str(err)}
    
################################################### END #####################################################
operations = {
    'Replace Strings':microsoft_word_replace_strings,
    'Replace Table':microsoft_word_replace_table
}
